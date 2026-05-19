from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("DSEasy — Project Overview")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Making visualisation of DSE-listed companies' performance easy.")
sr.italic = True
sr.font.size = Pt(11)

def h(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    if bold_prefix:
        b = p.add_run(bold_prefix)
        b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

h("Objective")
doc.add_paragraph(
    "DSEasy is a modern web dashboard that transforms raw Dar es Salaam Stock Exchange (DSE) "
    "market data into an interactive, investor-friendly interface — surfacing trends, liquidity, "
    "and momentum at a glance for individual investors and analysts."
).paragraph_format.space_after = Pt(2)

h("Stack")
bullet("React + TypeScript (strict), Vite, Chart.js via react-chartjs-2, Firebase Firestore.",
       bold_prefix="Frontend & Data: ")
bullet("Google Apps Script + Google Sheets pipeline feeding Firestore (trends/{symbol}/dailyClosingHistory).",
       bold_prefix="Ingestion: ")
bullet("Vercel.", bold_prefix="Hosting: ")

h("Core Features")
bullet("Market overview: gainers/losers, total volume, turnover, market cap.")
bullet("Ticker Trends: historical close prices, indicators, and toggleable overlays.")
bullet("Daily Glance Intel: at-a-glance market summary per scrape cycle.")
bullet("Sortable, detailed table view for granular company inspection.")

h("Roadmap")

h2 = doc.add_paragraph()
h2r = h2.add_run("Now — Financial Logic")
h2r.bold = True
h2r.font.size = Pt(10)
h2.paragraph_format.space_after = Pt(0)
bullet("Moving Averages (SMA/EMA 20/50) overlay on Price Action.")
bullet("RSI (14-day) indicator chart on Ticker Trends.")
bullet("VWAP calculation and display.")
bullet("Comparative performance: multi-symbol normalised % chart.")

h3 = doc.add_paragraph()
h3r = h3.add_run("Next — UX & Platform")
h3r.bold = True
h3r.font.size = Pt(10)
h3.paragraph_format.space_after = Pt(0)
bullet("Improved loading states & transitions.")
bullet("Real user authentication (retire placeholder profile).")
bullet("PWA + push notifications; functional settings menu.")
bullet("Consistent footer; dark-mode polish on Trends dropdown.")

h4 = doc.add_paragraph()
h4r = h4.add_run("Later — Intel & Infrastructure")
h4r.bold = True
h4r.font.size = Pt(10)
h4.paragraph_format.space_after = Pt(0)
bullet("Anomaly flags & Market Momentum Score (-100 → +100) with sparkline.")
bullet("Watchlist-aware personalised intel paragraphs.")
bullet("DSE & CMSA announcement scraping; AI-summarised overnight news.")
bullet("Service-layer refactor; unit tests for financial calculations; Level 2 insights.")

h("Working Principles")
bullet("Branch off latest master: feat/issue-N-..., fix/issue-N-..., chore/issue-N-...")
bullet("Never push to master. Ask before squash-merging.")
bullet("StockData field names are canonical; raw Firestore names live only in RawStockDoc.")
bullet("One mapping point: toStockData() in useMarketQuery.ts.")

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("ds-easy.vercel.app   •   Stay focused. Ship small. Measure the market.")
fr.italic = True
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

out = "/home/user/DSEasy/docs/DSEasy-Project-Overview.docx"
import os
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print("Wrote", out)
